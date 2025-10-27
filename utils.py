<<<<<<< HEAD
# utils.py
import docx2txt
import PyPDF2
from fpdf import FPDF
from docx import Document
import os
import time
import json # Added json import

# --- Existing Functions ---

def extract_text_from_file(file):
    """Extracts plain text from PDF, DOCX, or TXT file objects."""
    try:
        if file.name.endswith(".pdf"):
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text: # Append only if text is extracted
                     text += page_text + "\n" # Add newline between pages
            return text.strip() # Remove leading/trailing whitespace

        elif file.name.endswith(".docx"):
            # Use a temporary file path as docx2txt might work better with paths
            # Ensure temp file has unique name to avoid conflicts if run concurrently
            temp_path = f"temp_{int(time.time())}_{os.getpid()}.docx"
            try:
                with open(temp_path, "wb") as f:
                    f.write(file.getvalue())
                text = docx2txt.process(temp_path)
            finally:
                if os.path.exists(temp_path):
                    os.remove(temp_path) # Clean up temp file
            return text.strip()

        elif file.name.endswith(".txt"):
            # Try common encodings
            try:
                # Reset file pointer before reading again
                file.seek(0)
                return file.read().decode("utf-8").strip()
            except UnicodeDecodeError:
                try:
                    file.seek(0) # Reset file pointer
                    return file.read().decode("latin-1").strip()
                except Exception as e:
                     print(f"Error decoding .txt file: {e}")
                     return "" # Return empty on error
            except Exception as e:
                print(f"Error reading .txt file: {e}")
                return ""
    except Exception as e:
        # Catch potential errors like file corruption or unexpected types
        print(f"Error extracting text from {getattr(file, 'name', 'file')}: {e}")
        return "" # Return empty string on any error

    return "" # Fallback

def save_edited_resume(text, format="docx"):
    """Saves the provided text as a DOCX or PDF file in an 'uploads' directory."""
    # Ensure uploads directory exists relative to this script's location
    base_dir = os.path.dirname(os.path.abspath(__file__))
    upload_dir = os.path.join(base_dir, "uploads")
    os.makedirs(upload_dir, exist_ok=True)

    timestamp = int(time.time())
    file_name = f"updated_resume_{timestamp}.{format}"
    path = os.path.join(upload_dir, file_name)

    try:
        if format == "docx":
            doc = Document()
            # Ensure text is not None before adding
            doc.add_paragraph(str(text) if text is not None else "")
            doc.save(path)
        elif format == "pdf":
            pdf = FPDF()
            pdf.add_page()
            # Add a font that supports a wider range of characters (like DejaVu)
            # You might need to install DejaVu fonts or provide the path to the .ttf file
            try:
                # Try adding Unicode font (ensure ttf file path is correct or font is installed)
                # pdf.add_font("DejaVu", "", "DejaVuSans.ttf", uni=True)
                pdf.set_font("Arial", size=12) # Fallback to Arial if DejaVu fails
            except RuntimeError:
                print("Warning: DejaVu font not found or path incorrect. Using Arial (may have limited Unicode support).")
                pdf.set_font("Arial", size=12)

            # Use multi_cell for automatic line breaks, encode properly for FPDF
            pdf_text = str(text) if text is not None else ""
            # Encode using latin-1 with replacement for unsupported characters
            pdf.multi_cell(0, 10, pdf_text.encode('latin-1', 'replace').decode('latin-1'))
            pdf.output(path)
        else:
             print(f"Unsupported format: {format}")
             return None

        print(f"Saved edited resume to: {path}")
        return path
    except Exception as e:
        print(f"Error saving edited resume to {path}: {e}")
        return None

# --- Appended JSON Functions ---

def load_json(path):
    """Load JSON from path, return Python object or None on failure."""
    try:
        if not os.path.exists(path):
            print(f"Warning: JSON file not found at {path}")
            return None
        with open(path, "r", encoding="utf-8") as f:
            # Check if file is empty before loading
            content = f.read()
            if not content:
                print(f"Warning: JSON file is empty at {path}")
                return None
            data = json.loads(content)
            return data
    except json.JSONDecodeError as e:
        print(f"Error decoding JSON from {path}: {e}")
        return None
    except Exception as e:
        print(f"Error loading JSON {path}: {e}")
        return None

def save_json(path, obj):
    """Save Python object as JSON (indent=4, ensure_ascii=False)."""
    try:
        # Ensure the directory exists before writing
        output_dir = os.path.dirname(path)
        if output_dir: # Check if dirname returned something (not empty for root files)
             os.makedirs(output_dir, exist_ok=True)

        with open(path, "w", encoding="utf-8") as f:
            # Ensure obj is not None before dumping, default to empty list if None
            json.dump(obj if obj is not None else [], f, ensure_ascii=False, indent=4)
        print(f"Successfully saved JSON to {path}")
        return True
    except TypeError as e:
        print(f"Error saving JSON to {path}: Object might not be JSON serializable. {e}")
        return False
    except Exception as e:
        print(f"Error saving JSON {path}: {e}")
        return False
=======
import docx2txt
import PyPDF2
from fpdf import FPDF
from docx import Document
import os
import time

def extract_text_from_file(file):
    if file.name.endswith(".pdf"):
        reader = PyPDF2.PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text

    elif file.name.endswith(".docx"):
        return docx2txt.process(file)

    elif file.name.endswith(".txt"):
        return file.read().decode("utf-8")

    return ""

def save_edited_resume(text, format="docx"):
    os.makedirs("uploads", exist_ok=True)
    timestamp = int(time.time())
    path = f"uploads/updated_resume_{timestamp}.{format}"

    if format == "docx":
        doc = Document()
        doc.add_paragraph(text)
        doc.save(path)
    else:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, text)
        pdf.output(path)

    return path
>>>>>>> 1513229e212edae69176b12f7d1ef905a471d922
